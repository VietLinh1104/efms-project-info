from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = "output/Ket_luan_he_thong_EFMS.docx"
FONT_NAME = "Times New Roman"
FONT_SIZE = 13


def set_run_font(run, size=FONT_SIZE, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)


def configure_style(style, size=FONT_SIZE, bold=False):
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.get_or_add_rPr()
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=11)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph(style="Body Text EFMS")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.27)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_section_heading(doc, number, title):
    paragraph = doc.add_paragraph(style="Section Heading EFMS")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(f"{number}.\t{title}")
    set_run_font(run, bold=True)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="EFMS Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    configure_style(normal)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    body_style = doc.styles.add_style("Body Text EFMS", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(body_style)

    heading_style = doc.styles.add_style("Section Heading EFMS", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(heading_style, bold=True)

    bullet_style = doc.styles.add_style("EFMS Bullet", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(bullet_style)
    bullet_style.base_style = normal
    bullet_style._element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.extend([ilvl, num_id])
    bullet_style._element.pPr.append(num_pr)

    footer = section.footer
    add_page_number(footer.paragraphs[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run("KẾT LUẬN")
    set_run_font(title_run, size=14, bold=True)

    add_section_heading(doc, "1", "Kết quả đạt được")

    add_body_paragraph(
        doc,
        "Sau quá trình nghiên cứu, phân tích, thiết kế và xây dựng, đồ án đã hoàn thành hệ thống EFMS "
        "(Enterprise Financial Management System) với các chức năng cơ bản đáp ứng nhu cầu quản lý tài chính và "
        "kế toán trong doanh nghiệp."
    )

    add_body_paragraph(
        doc,
        "Hệ thống đã xây dựng thành công các chức năng quản lý khách hàng, nhà cung cấp, tài khoản kế toán, tài "
        "khoản ngân hàng, hóa đơn phải thu, hóa đơn phải trả và thanh toán. Người dùng có thể tạo, cập nhật, xác "
        "nhận hoặc hủy hóa đơn; ghi nhận thu, chi tiền; phân bổ thanh toán và theo dõi công nợ."
    )

    add_body_paragraph(
        doc,
        "Đối với quản lý tài chính và kế toán trưởng, EFMS hỗ trợ kiểm tra, phê duyệt hoặc từ chối hóa đơn mua "
        "hàng, đồng thời cung cấp Dashboard theo dõi các khoản phải thu, phải trả, thanh toán trong kỳ và hóa đơn "
        "đang chờ duyệt. Hệ thống cũng tự động sinh các bút toán kế toán từ nghiệp vụ thanh toán."
    )

    add_body_paragraph(
        doc,
        "Đối với quản trị viên, hệ thống hỗ trợ quản lý doanh nghiệp, người dùng, vai trò và quyền truy cập theo mô "
        "hình RBAC. Cơ chế đa doanh nghiệp giúp phân tách dữ liệu và bảo đảm người dùng chỉ được truy cập thông tin "
        "thuộc công ty của mình."
    )

    add_body_paragraph(
        doc,
        "Bên cạnh đó, hệ thống đã tích hợp Claude AI thông qua MCP Server nhằm hỗ trợ tra cứu dữ liệu bằng ngôn ngữ "
        "tự nhiên, phân tích rủi ro hóa đơn và cung cấp gợi ý cho người phê duyệt. AI đóng vai trò hỗ trợ, còn quyết "
        "định nghiệp vụ cuối cùng vẫn thuộc về người dùng có thẩm quyền."
    )

    add_body_paragraph(
        doc,
        "Hệ thống được xây dựng theo kiến trúc microservices với React, Java 21, Spring Boot, Spring Cloud và "
        "PostgreSQL. API Gateway thực hiện định tuyến và xác thực JWT, góp phần bảo đảm tính bảo mật, khả năng bảo "
        "trì và mở rộng hệ thống."
    )

    add_section_heading(doc, "2", "Hạn chế")

    add_body_paragraph(
        doc,
        "Mặc dù đã hoàn thành các chức năng chính trong phạm vi đồ án, hệ thống vẫn còn một số hạn chế cần tiếp tục "
        "nghiên cứu và hoàn thiện:"
    )

    limitations = [
        "Giao diện chưa được tối ưu hoàn toàn cho thiết bị di động.",
        "Hệ thống chưa triển khai đầy đủ các nghiệp vụ kế toán chuyên sâu như kỳ kế toán, khóa sổ, báo cáo tài "
        "chính và đối soát ngân hàng.",
        "Claude AI hiện chủ yếu hỗ trợ tra cứu và phân tích cơ bản, chưa có khả năng học theo lịch sử hoạt động của "
        "từng doanh nghiệp.",
        "Chưa tích hợp xác thực OTP, xác thực đa yếu tố, ngân hàng, hóa đơn điện tử và các hệ thống ERP bên ngoài.",
        "Các giải pháp tối ưu hiệu suất, giám sát, sao lưu và bảo mật nâng cao cho môi trường triển khai quy mô lớn "
        "chưa được áp dụng đầy đủ.",
    ]
    for item in limitations:
        add_bullet(doc, item)

    add_body_paragraph(
        doc,
        "Do thời gian thực hiện và kinh nghiệm triển khai thực tế còn hạn chế nên hệ thống vẫn còn một số thiếu sót "
        "cần tiếp tục hoàn thiện trong tương lai."
    )

    add_section_heading(doc, "3", "Hướng phát triển đồ án")

    add_body_paragraph(
        doc,
        "Trong thời gian tới, hệ thống EFMS có thể được tiếp tục phát triển và mở rộng theo các hướng sau:"
    )

    directions = [
        "Tối ưu giao diện responsive trên điện thoại và máy tính bảng.",
        "Bổ sung các nghiệp vụ kế toán chuyên sâu, báo cáo tài chính và đối soát ngân hàng.",
        "Phát triển Claude AI để hỗ trợ dự báo dòng tiền, phát hiện bất thường và đánh giá rủi ro công nợ.",
        "Xây dựng quy trình phê duyệt nhiều cấp linh hoạt theo hạn mức và chính sách của từng doanh nghiệp.",
        "Hoàn thiện chức năng OTP, xác thực đa yếu tố và tăng cường bảo mật hệ thống.",
        "Tích hợp ngân hàng, hóa đơn điện tử, cổng thanh toán và các hệ thống ERP.",
        "Tối ưu hiệu suất, khả năng chịu tải, giám sát, sao lưu dữ liệu và phát triển ứng dụng mobile.",
    ]
    for item in directions:
        add_bullet(doc, item)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
