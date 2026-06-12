from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = "output/Tai_lieu_tham_khao_he_thong_EFMS.docx"
FONT_NAME = "Times New Roman"
ACCESS_DATE = "12/06/2026"


def set_run_font(run, size=13, bold=False, italic=False, color=None, underline=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        run._element.rPr.rFonts.set(qn(f"w:{key}"), FONT_NAME)


def configure_style(style, size=13):
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        style._element.rPr.rFonts.set(qn(f"w:{key}"), FONT_NAME)


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT_NAME)
    run_properties.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "26")
    run_properties.append(size)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "26")
    run_properties.append(size_cs)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([run_properties, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_reference(doc, number, author, title, url=None, note=None):
    paragraph = doc.add_paragraph(style="Reference EFMS")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.first_line_indent = Cm(-1.0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    prefix = paragraph.add_run(f"[{number}] {author}, ")
    set_run_font(prefix)
    title_run = paragraph.add_run(title)
    set_run_font(title_run, italic=True)

    if note:
        note_run = paragraph.add_run(f", {note}")
        set_run_font(note_run)

    if url:
        separator = paragraph.add_run(". Địa chỉ: ")
        set_run_font(separator)
        add_hyperlink(paragraph, url, url)
        accessed = paragraph.add_run(f" (truy cập ngày {ACCESS_DATE}).")
        set_run_font(accessed)
    else:
        period = paragraph.add_run(".")
        set_run_font(period)


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    configure_style(document.styles["Normal"])
    reference_style = document.styles.add_style("Reference EFMS", WD_STYLE_TYPE.PARAGRAPH)
    configure_style(reference_style)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run("TÀI LIỆU THAM KHẢO")
    set_run_font(title_run, size=14, bold=True)

    references = [
        (
            "Nhóm phát triển EFMS",
            "Tài liệu kiến trúc hệ thống EFMS",
            None,
            "tài liệu nội bộ dự án",
        ),
        (
            "Nhóm phát triển EFMS",
            "Tài liệu mô tả luồng nghiệp vụ và phân quyền hệ thống EFMS",
            None,
            "tài liệu nội bộ dự án",
        ),
        (
            "Oracle",
            "Java Platform, Standard Edition 21 Documentation",
            "https://docs.oracle.com/en/java/javase/21/",
            None,
        ),
        (
            "Spring",
            "Spring Boot 3.3 Reference Documentation",
            "https://docs.spring.io/spring-boot/3.3/index.html",
            None,
        ),
        (
            "Spring",
            "Spring Cloud Gateway Reference Documentation",
            "https://docs.spring.io/spring-cloud-gateway/reference/index.html",
            None,
        ),
        (
            "Spring",
            "Spring Security Reference",
            "https://docs.spring.io/spring-security/reference/index.html",
            None,
        ),
        (
            "The PostgreSQL Global Development Group",
            "PostgreSQL Documentation",
            "https://www.postgresql.org/docs/current/",
            None,
        ),
        (
            "Meta Open Source",
            "React Reference",
            "https://react.dev/reference/react",
            None,
        ),
        (
            "Apache Software Foundation",
            "Apache Maven Documentation",
            "https://maven.apache.org/guides/index.html",
            None,
        ),
        (
            "M. Jones, J. Bradley và N. Sakimura",
            "RFC 7519: JSON Web Token (JWT)",
            "https://www.rfc-editor.org/rfc/rfc7519.html",
            "Internet Engineering Task Force, 2015",
        ),
        (
            "D. Hardt",
            "RFC 6749: The OAuth 2.0 Authorization Framework",
            "https://www.rfc-editor.org/rfc/rfc6749.html",
            "Internet Engineering Task Force, 2012",
        ),
        (
            "Model Context Protocol",
            "Model Context Protocol Specification",
            "https://modelcontextprotocol.io/specification/2025-11-25",
            None,
        ),
        (
            "OWASP Foundation",
            "OWASP Top 10: Web Application Security Risks",
            "https://owasp.org/www-project-top-ten/",
            None,
        ),
    ]

    for index, (author, title_text, url, note) in enumerate(references, start=1):
        add_reference(document, index, author, title_text, url, note)

    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
